import os
import re
os.environ['HUGGINGFACE_HUB_OFFLINE'] = '1'  # Try offline mode first

from flask import Flask, request, jsonify
from flask_cors import CORS
import PyPDF2
from difflib import SequenceMatcher
import glob
import docx
import json
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from io import BytesIO
from datetime import datetime

# Note: Turnitin-style detection doesn't require sentence transformers
# We use n-gram matching and sequence matching instead
SentenceTransformer = None
util = None
SentenceTransformer_available = False

app = Flask(__name__)
CORS(app)

# Define the base path for thesis files (relative to the Laravel/PHP project root)
# The Flask app is in /fyp/web-test, so thesis files are in ../uploads/thesis/
BASE_UPLOAD_PATH = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'uploads', 'thesis'))

print("=" * 60)
print("TURNITIN-STYLE PLAGIARISM DETECTION API")
print("=" * 60)
print(f"Detection method: N-gram matching + Consecutive text matching")
print(f"Weights: 70% Consecutive match + 30% Phrase match")
print(f"Upload path: {BASE_UPLOAD_PATH}")
print("=" * 60)

# Function to extract text from uploaded file
def extract_text(file):
    filename = file.filename.lower()
    text = ""

    try:
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif filename.endswith('.txt'):
            text = file.read().decode('utf-8', errors='ignore')
        elif filename.endswith('.docx'):
            # For docx files, use python-docx if available
            doc = docx.Document(file)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            text = ""
    except Exception as e:
        print(f"Error extracting text from uploaded file: {e}")
        text = ""
    
    return text.strip()

# Function to extract text from file path
def extract_text_from_path(file_path):
    text = ""
    filename = file_path.lower()
    
    # Resolve the file path - it may be relative or absolute
    if not os.path.isabs(file_path):
        # If relative, resolve it relative to the base upload path
        resolved_path = os.path.normpath(os.path.join(BASE_UPLOAD_PATH, os.path.basename(file_path)))
    else:
        resolved_path = os.path.normpath(file_path)
    
    print(f"Attempting to read file: {file_path}")
    print(f"Resolved path: {resolved_path}")
    
    # Verify the file exists
    if not os.path.exists(resolved_path):
        print(f"File not found at resolved path: {resolved_path}")
        # Try alternate paths
        if os.path.exists(file_path):
            resolved_path = file_path
            print(f"Found at original path: {file_path}")
        else:
            print(f"File not found at any path")
            return ""
    
    try:
        if filename.endswith('.pdf'):
            with open(resolved_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif filename.endswith('.txt'):
            with open(resolved_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            try:
                doc = docx.Document(resolved_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as docx_error:
                print(f"Error reading DOCX file: {docx_error}")
                # For .doc files, return empty text as python-docx doesn't support old Word format
                text = ""
    except Exception as e:
        print(f"Error reading file {resolved_path}: {e}")
        text = ""
    
    return text.strip()

def clean_text(text):
    """Clean and normalize text for comparison"""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters and convert to lowercase
    text = re.sub(r'[^\w\s]', '', text.lower())
    return text.strip()


def is_exact_match(text1, text2):
    """Return True only when the two texts are 100% identical after normalization.

    Normalization includes removing extra whitespace, special characters, and
    converting to lowercase. This triggers an exact match when the cleaned
    versions are identical.
    """
    if text1 is None or text2 is None:
        return False
    return clean_text(text1) == clean_text(text2)


def extract_sentences(text):
    """Extract sentences from text, split by common sentence delimiters."""
    if not text:
        return []
    # Split by periods, question marks, exclamation marks, and newlines
    sentences = re.split(r'(?<=[.!?])\s+|(?<=\n)\s*', text.strip())
    # Filter out empty sentences and strip whitespace
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


def extract_ngrams(text, n=5):
    """Extract n-grams (word sequences) from text for phrase-level plagiarism detection.
    Like Turnitin, this detects matching phrases/word sequences.
    """
    words = clean_text(text).split()
    ngrams = []
    for i in range(len(words) - n + 1):
        ngram = ' '.join(words[i:i+n])
        if len(ngram) > 10:  # Only consider meaningful n-grams
            ngrams.append(ngram)
    return ngrams


def find_matching_ngrams(uploaded_text, existing_text, ngram_size=5):
    """Find matching n-grams between two documents.
    Like Turnitin, this detects consecutive phrase matches.
    Returns count of matching n-grams and percentage of document covered.
    """
    uploaded_ngrams = set(extract_ngrams(uploaded_text, ngram_size))
    existing_ngrams = set(extract_ngrams(existing_text, ngram_size))
    
    if not uploaded_ngrams or not existing_ngrams:
        return 0, 0.0
    
    matching_ngrams = uploaded_ngrams.intersection(existing_ngrams)
    coverage_percentage = (len(matching_ngrams) / len(uploaded_ngrams)) * 100 if uploaded_ngrams else 0
    
    return len(matching_ngrams), coverage_percentage


def find_longest_common_substring(text1, text2, min_length=20):
    """Find longest common substrings between texts.
    Like Turnitin, detects consecutive identical phrases.
    """
    clean1 = clean_text(text1)
    clean2 = clean_text(text2)
    
    if not clean1 or not clean2:
        return []
    
    matches = []
    
    # Use SequenceMatcher to find matching blocks
    matcher = SequenceMatcher(None, clean1, clean2)
    
    for block in matcher.get_matching_blocks():
        if block.size >= min_length:
            # Extract the actual matched text (without context padding)
            # This ensures each match shows distinct content
            start1 = block.a
            end1 = min(len(clean1), block.a + block.size)
            start2 = block.b
            end2 = min(len(clean2), block.b + block.size)
            
            match_text_uploaded = clean1[start1:end1]
            match_text_existing = clean2[start2:end2]
            
            matches.append({
                'uploaded': match_text_uploaded[:200],  # Show more of the actual match
                'existing': match_text_existing[:200],  # Show more of the actual match
                'match_length': block.size
            })
    
    return sorted(matches, key=lambda x: x['match_length'], reverse=True)


def find_similar_sentences(uploaded_text, existing_text, threshold=0.65):
    """Find similar sequences and phrases like Turnitin plagiarism detection.
    
    Turnitin-style detection focuses on:
    - Matching word sequences (phrases)
    - Consecutive text matches
    - Document coverage percentage
    
    Returns a list of flagged similarities with type and match data.
    """
    similar_pairs = []
    
    if not uploaded_text or not existing_text:
        return similar_pairs
    
    try:
        # 1. Find matching n-grams (word sequences/phrases)
        matching_count, coverage = find_matching_ngrams(uploaded_text, existing_text, ngram_size=5)
        
        # Flag PHRASE_MATCH if coverage >= 2% (lowered from 65%)
        if coverage >= 2.0:
            similar_pairs.append({
                'type': 'PHRASE_MATCH',
                'uploaded': uploaded_text[:300],
                'existing': existing_text[:300],
                'similarity': round(coverage, 2),
                'match_count': matching_count,
                'description': f'Detected {matching_count} matching phrases covering {round(coverage, 2)}% of uploaded document'
            })
        
        # 2. Find longest common consecutive sequences
        common_sequences = find_longest_common_substring(uploaded_text, existing_text, min_length=20)
        
        for sequence in common_sequences:  # Show ALL sequences, not just top 5
            sequence_coverage = (sequence['match_length'] / len(clean_text(uploaded_text))) * 100 if uploaded_text else 0
            # Flag CONSECUTIVE_MATCH if coverage >= 1% (lowered from 32.5%)
            if sequence_coverage >= 1.0:
                similar_pairs.append({
                    'type': 'CONSECUTIVE_MATCH',
                    'uploaded': sequence['uploaded'],
                    'existing': sequence['existing'],
                    'similarity': round(sequence_coverage, 2),
                    'match_length': sequence['match_length'],
                    'description': f'Found {sequence["match_length"]} consecutive matching characters'
                })
        
    except Exception as e:
        print(f"Error finding similar phrases (Turnitin-style): {e}")
    
    return similar_pairs

def calculate_text_similarity(text1, text2):
    """Calculate similarity using Turnitin-style plagiarism detection.
    
    Like Turnitin, this prioritizes:
    1. Consecutive matching text (word-for-word plagiarism) - 70%
    2. Phrase/n-gram matching (paraphrased plagiarism) - 30%
    
    Returns a percentage indicating likelihood of plagiarism.
    """
    clean_text1 = clean_text(text1)
    clean_text2 = clean_text(text2)
    
    if not clean_text1 or not clean_text2:
        return 0
    
    try:
        # 1. CONSECUTIVE MATCH DETECTION (Word-for-word plagiarism) - 70% weight
        # SequenceMatcher finds longest common substrings
        matcher = SequenceMatcher(None, clean_text1, clean_text2)
        matching_blocks = matcher.get_matching_blocks()
        
        # Calculate coverage of matching blocks
        total_match_length = sum(block.size for block in matching_blocks)
        consecutive_similarity = (total_match_length / len(clean_text1)) if clean_text1 else 0
        
        # 2. PHRASE MATCHING (Paraphrased plagiarism) - 30% weight
        # Check how many n-grams (word sequences) match
        matching_ngrams, phrase_coverage = find_matching_ngrams(text1, text2, ngram_size=5)
        phrase_similarity = phrase_coverage / 100.0  # Convert percentage to decimal
        
        # 3. WORD OVERLAP (Semantic similarity)
        # Common words between documents
        words1 = set(clean_text1.split())
        words2 = set(clean_text2.split())
        
        if words1 and words2:
            intersection = len(words1.intersection(words2))
            word_overlap = intersection / len(words1) if words1 else 0
        else:
            word_overlap = 0
        
        # TURNITIN-STYLE CALCULATION:
        # 70% weight on consecutive matches (word-for-word plagiarism)
        # 30% weight on phrase matches (paraphrased content)
        turnitin_similarity = (consecutive_similarity * 0.70) + (phrase_similarity * 0.30)
        
        # Cap at 100%
        turnitin_similarity = min(turnitin_similarity, 1.0)
        
        print(f"[SIMILARITY BREAKDOWN]")
        print(f"  Consecutive match: {round(consecutive_similarity*100, 2)}%")
        print(f"  Phrase match: {round(phrase_similarity*100, 2)}%")
        print(f"  Word overlap: {round(word_overlap*100, 2)}%")
        print(f"  TURNITIN SCORE: {round(turnitin_similarity*100, 2)}%")
        
        return turnitin_similarity
        
    except Exception as e:
        print(f"Error in Turnitin-style similarity calculation: {e}")
        # Fallback to basic consecutive match
        return SequenceMatcher(None, clean_text1, clean_text2).ratio()

@app.route('/api/check_similarity', methods=['POST'])
def check_similarity():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        uploaded_file = request.files['file']
        if uploaded_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get existing files data from PHP
        existing_files_json = request.form.get('existing_files', '[]')
        try:
            existing_files = json.loads(existing_files_json)
        except json.JSONDecodeError:
            existing_files = []
        
        print(f"\n=== SIMILARITY CHECK START ===")
        print(f"Uploaded file: {uploaded_file.filename}")
        print(f"Number of existing files to compare: {len(existing_files)}")
        
        # Extract text from uploaded file
        uploaded_text = extract_text(uploaded_file)
        print(f"Uploaded text length: {len(uploaded_text)} characters")
        print(f"First 100 chars of uploaded: {uploaded_text[:100]}")
        
        if not uploaded_text:
            return jsonify({'error': 'Could not extract text from uploaded file'}), 400
        
        if not existing_files:
            return jsonify({
                'max_similarity': 0,
                'message': '✅ No existing thesis files found to compare against. You can submit your thesis.',
                'similar_files_list': [],
                'can_submit': True
            })
        
        max_similarity = 0
        similar_file = None
        exact_match = False
        exact_match_file = None
        similar_sentences = []
        all_matches = []  # Store ALL matches with similarity > 0%
        
        # Check similarity against each existing file
        for file_info in existing_files:
            file_path = file_info['path']
            filename = file_info['filename']
            
            print(f"\nProcessing existing file: {filename}")
            print(f"File path from PHP: {file_path}")
            existing_text = extract_text_from_path(file_path)
            print(f"Existing text length: {len(existing_text)} characters")
            print(f"First 100 chars of existing: {existing_text[:100]}")
            
            if existing_text:
                # Check exact match first
                if is_exact_match(uploaded_text, existing_text):
                    exact_match = True
                    exact_match_file = filename
                    max_similarity = 100.0
                    similar_file = filename
                    print(f"✓ EXACT MATCH found with {filename}")
                    
                    # Find similar sentences even for exact match
                    similar_sentences = find_similar_sentences(uploaded_text, existing_text, threshold=0.65)
                    print(f"Plagiarism flags detected: {len(similar_sentences)}")
                    
                    # Add to matches list
                    all_matches.append({
                        'filename': filename,
                        'similarity': 100.0,
                        'similar_sentences': similar_sentences
                    })
                    break

                # Calculate Turnitin-style plagiarism similarity
                similarity = calculate_text_similarity(uploaded_text, existing_text)
                similarity_percentage = round(similarity * 100, 2)

                print(f"\n>>> PLAGIARISM CHECK vs {filename}: {similarity_percentage}%")

                # Add to matches list if similarity > 0%
                if similarity_percentage > 0:
                    # Find plagiarized phrases and consecutive matches
                    file_similar_sentences = find_similar_sentences(uploaded_text, existing_text, threshold=0.65)
                    
                    all_matches.append({
                        'filename': filename,
                        'similarity': similarity_percentage,
                        'similar_sentences': file_similar_sentences
                    })
                    
                    print(f"Added to matches: {filename} - {similarity_percentage}%")

                if similarity_percentage > max_similarity:
                    max_similarity = similarity_percentage
                    similar_file = filename
                    similar_sentences = file_similar_sentences
            else:
                print(f"✗ Could not extract text from {filename}")
        
        # Sort matches by similarity in descending order
        all_matches.sort(key=lambda x: x['similarity'], reverse=True)
        
        # Create message
        if len(all_matches) == 0:
            message = "✅ Similarity check completed! No matching documents found."
        else:
            message = f"✅ Similarity check completed! Found {len(all_matches)} matching document(s)."

        response_data = {
            'max_similarity': max_similarity,
            'message': message,
            'similar_file': similar_file,
            'exact_match': exact_match,
            'exact_match_file': exact_match_file,
            'similar_sentences': similar_sentences,
            'similar_files_list': all_matches,  # List of all matches
            'can_submit': True
        }
        
        print(f"=== RETURNING RESPONSE ===")
        print(f"Found {len(all_matches)} matches")
        for match in all_matches:
            print(f"  - {match['filename']}: {match['similarity']}%")
        print(f"=== SIMILARITY CHECK END ===\n")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Error in check_similarity: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/health', methods=['GET'])
def health_check():
    # Check if upload directory exists
    upload_dir_exists = os.path.exists(BASE_UPLOAD_PATH)
    return jsonify({
        'status': 'ok', 
        'message': 'Similarity checker API is running',
        'upload_path': BASE_UPLOAD_PATH,
        'upload_dir_exists': upload_dir_exists
    })


@app.route('/api/generate_report', methods=['POST'])
def generate_report():
    """Generate a PDF report of all matching documents and their similar sentences."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        uploaded_filename = data.get('uploaded_filename', 'Unknown')
        similar_files_list = data.get('similar_files_list', [])
        
        # Create PDF in memory
        pdf_buffer = BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        
        # Container for PDF elements
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e3c72'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2a5298'),
            spaceAfter=12
        )
        
        document_heading_style = ParagraphStyle(
            'DocumentHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#34568B'),
            spaceAfter=8
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8
        )
        
        # Title
        elements.append(Paragraph("Thesis Similarity Report", title_style))
        elements.append(Spacer(1, 0.3 * inch))
        
        # Report metadata
        metadata = [
            ['Report Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Uploaded File:', uploaded_filename],
            ['Total Matching Documents:', str(len(similar_files_list))]
        ]
        
        metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        
        elements.append(metadata_table)
        elements.append(Spacer(1, 0.3 * inch))
        
        # Check if there are any matches
        if not similar_files_list or len(similar_files_list) == 0:
            elements.append(Paragraph("Summary", heading_style))
            elements.append(Paragraph("✅ No matching documents found. Your thesis is unique!", normal_style))
        else:
            # Summary table of all matches
            elements.append(Paragraph("Summary", heading_style))
            elements.append(Spacer(1, 0.1 * inch))
            
            summary_data = [['Document Name', 'Similarity %', 'Similar Sentences']]
            for match in similar_files_list:
                similar_count = len(match.get('similar_sentences', []))
                summary_data.append([
                    match['filename'],
                    f"{match['similarity']}%",
                    str(similar_count)
                ])
            
            summary_table = Table(summary_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34568B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            elements.append(summary_table)
            elements.append(Spacer(1, 0.3 * inch))
            
            # Detailed report for each matching document
            elements.append(Paragraph("Detailed Matches", heading_style))
            elements.append(Spacer(1, 0.15 * inch))
            
            for doc_idx, match in enumerate(similar_files_list, 1):
                # Document heading
                doc_heading = f"{doc_idx}. {match['filename']} (Similarity: {match['similarity']}%)"
                elements.append(Paragraph(doc_heading, document_heading_style))
                
                similar_sentences = match.get('similar_sentences', [])
                
                if similar_sentences and len(similar_sentences) > 0:
                    # Add similar sentences for this document
                    for sent_idx, sentence in enumerate(similar_sentences, 1):
                        # Sentence match header
                        sent_header = f"Match #{sent_idx} - Similarity: {sentence.get('similarity', 0)}%"
                        sent_style = ParagraphStyle(
                            f'SentenceHeader{doc_idx}{sent_idx}',
                            parent=styles['Normal'],
                            fontSize=10,
                            textColor=colors.HexColor('#2a5298'),
                            fontName='Helvetica-Bold',
                            spaceAfter=6
                        )
                        elements.append(Paragraph(sent_header, sent_style))
                        
                        # Comparison table for this sentence
                        comparison_data = [
                            ['Source', 'Text'],
                            ['Your Document', sentence.get('uploaded', 'N/A')],
                            ['Existing Document', sentence.get('existing', 'N/A')]
                        ]
                        
                        comparison_table = Table(comparison_data, colWidths=[1.2*inch, 5.3*inch])
                        comparison_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34568B')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 10),
                            ('FONTSIZE', (0, 1), (-1, -1), 9),
                            ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#E8F4F8')),
                            ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#FFF8E8')),
                            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                            ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ]))
                        
                        elements.append(comparison_table)
                        elements.append(Spacer(1, 0.15 * inch))
                else:
                    # No similar sentences for this document
                    elements.append(Paragraph("No detailed similar sentences detected.", normal_style))
                    elements.append(Spacer(1, 0.15 * inch))
                
                # Page break after each document (except the last one)
                if doc_idx < len(similar_files_list):
                    elements.append(PageBreak())
        
        # Footer
        elements.append(Spacer(1, 0.3 * inch))
        elements.append(Paragraph(
            "<i>This report was automatically generated by the USIM eThesis Similarity Checker.</i>",
            normal_style
        ))
        
        # Build PDF
        doc.build(elements)
        
        # Return PDF as response
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue(), 200, {
            'Content-Type': 'application/pdf',
            'Content-Disposition': 'inline; filename="similarity_report.pdf"'
        }
        
    except Exception as e:
        print(f"Error generating report: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500

if __name__ == '__main__':
    print("Starting AI-Powered Thesis Similarity Checker API...")
    print("API will be available at: http://localhost:5000")
    print("Features: Simple percentage-based similarity analysis")
    app.run(debug=True, host='0.0.0.0', port=5000)